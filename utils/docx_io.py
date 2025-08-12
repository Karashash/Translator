from docx.text.paragraph import Paragraph
from docx.table import _Cell
from typing import Iterable, Tuple, List

def iter_all_paragraphs(doc) -> Iterable[Paragraph]:
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                yield p
        if section.footer:
            for p in section.footer.paragraphs:
                yield p

def iter_all_tables(doc) -> Iterable[_Cell]:
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield cell
    for section in doc.sections:
        if section.header:
            for tbl in section.header.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        yield cell
        if section.footer:
            for tbl in section.footer.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        yield cell

def replace_paragraph_text_from_spans(par: Paragraph, spans: List[Tuple[int, str]]):
    orig_runs = list(par.runs)
    # remove all runs
    for i in range(len(par.runs)-1, -1, -1):
        r = par.runs[i]
        r.clear(); r.text = ""
        r._element.getparent().remove(r._element)

    def clone(dst, src):
        dst.bold = src.bold
        dst.italic = src.italic
        dst.underline = src.underline
        dst.style = src.style
        try:
            dst.font.name = src.font.name
            dst.font.size = src.font.size
            dst.font.color.rgb = getattr(src.font.color, "rgb", None)
        except Exception:
            pass

    for rid, txt in spans:
        if rid is not None and rid < len(orig_runs):
            new_run = par.add_run(txt)
            clone(new_run, orig_runs[rid])
        else:
            par.add_run(txt)

def replace_cell_text_from_spans(cell: _Cell, spans: List[Tuple[int, str]]):
    text = "".join([t for _, t in spans])
    lines = text.split("\n")
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.add_paragraph() if i > 0 else cell.paragraphs[0]
        p.text = line
